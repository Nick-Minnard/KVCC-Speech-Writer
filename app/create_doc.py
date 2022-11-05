from docx import Document
from docx.shared import RGBColor, Pt

def create_speech_doc(speech):
  """Create speech document function"""
  d = Document()
  d.add_heading(speech.title, 0)
  current_section = 1
  # Create a styled word doc with docx and return it to be downloaded
  for point in speech.data["points"]:
    if point["section"] > current_section:
      current_section += 1
      r = d.add_paragraph().add_run("__________________________________________________________________________________________________________")
      r.font.color.rgb = RGBColor(77, 168, 208)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Pt(20 * (point["level"] - 1))
    r = p.add_run(point["prefix"] + point["name"] + ": ")
    r.bold = True
    r.font.size = Pt(13)
    r = p.add_run(point["value"])
    r.font.size = Pt(13)
  return d
  